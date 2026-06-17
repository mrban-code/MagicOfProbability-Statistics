from __future__ import annotations

import csv
import math
import os
import warnings
from dataclasses import dataclass
from itertools import product
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

warnings.simplefilter("ignore", pd.errors.PerformanceWarning)


ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = ROOT.parent
SOURCE_DATA = PROJECT_ROOT / "变量" / "analysis_dataset.csv"

DATA_DIR = ROOT / "data"
RESULTS_DIR = ROOT / "results"
FIG_DIR = ROOT / "figures"
DOCX_PATH = ROOT / "PM25_prediction_model_report.docx"

TEST_START = pd.Timestamp("2025-05-22")
ALPHAS = [0.0, 0.01, 0.1, 1.0, 10.0, 30.0, 100.0, 300.0, 1000.0]
RNG_SEED = 20260613
FOREST_CONFIGS = [
    {"n_trees": 90, "max_depth": 6, "min_leaf": 8, "max_features": "sqrt", "n_thresholds": 12},
    {"n_trees": 110, "max_depth": 8, "min_leaf": 6, "max_features": "sqrt", "n_thresholds": 12},
    {"n_trees": 120, "max_depth": 7, "min_leaf": 10, "max_features": 0.45, "n_thresholds": 10},
]
KNN_K_VALUES = [5, 10, 20, 35, 50, 80]
ENSEMBLE_WEIGHT_STEP = 0.05


@dataclass
class RidgeModel:
    feature_names: list[str]
    means: np.ndarray
    scales: np.ndarray
    beta: np.ndarray
    alpha: float

    def predict(self, frame: pd.DataFrame) -> np.ndarray:
        x = frame[self.feature_names].astype(float).to_numpy()
        x_std = (x - self.means) / self.scales
        xb = np.column_stack([np.ones(len(x_std)), x_std])
        return np.maximum(0.0, xb @ self.beta)

    def coefficient_frame(self, model_name: str, top_n: int = 30) -> pd.DataFrame:
        coef = pd.DataFrame(
            {
                "model": model_name,
                "feature": self.feature_names,
                "standardized_coefficient": self.beta[1:],
                "abs_standardized_coefficient": np.abs(self.beta[1:]),
            }
        )
        return coef.sort_values("abs_standardized_coefficient", ascending=False).head(top_n)


@dataclass
class TreeNode:
    value: float
    feature_idx: int | None = None
    threshold: float | None = None
    left: "TreeNode | None" = None
    right: "TreeNode | None" = None


@dataclass
class RandomForestModel:
    feature_names: list[str]
    means: np.ndarray
    scales: np.ndarray
    trees: list[TreeNode]
    params: dict[str, float | int | str]

    def predict(self, frame: pd.DataFrame) -> np.ndarray:
        x = frame[self.feature_names].astype(float).to_numpy()
        x_std = (x - self.means) / self.scales
        all_pred = np.column_stack([predict_tree(tree, x_std) for tree in self.trees])
        return np.maximum(0.0, all_pred.mean(axis=1))


@dataclass
class KNNAnalogModel:
    feature_names: list[str]
    means: np.ndarray
    scales: np.ndarray
    x_train: np.ndarray
    y_train: np.ndarray
    k: int

    def predict(self, frame: pd.DataFrame) -> np.ndarray:
        x = frame[self.feature_names].astype(float).to_numpy()
        x_std = (x - self.means) / self.scales
        preds = []
        for start in range(0, len(x_std), 128):
            chunk = x_std[start : start + 128]
            d2 = ((chunk[:, None, :] - self.x_train[None, :, :]) ** 2).mean(axis=2)
            idx = np.argpartition(d2, kth=min(self.k, len(self.y_train) - 1), axis=1)[:, : self.k]
            nearest_d = np.take_along_axis(d2, idx, axis=1)
            nearest_y = self.y_train[idx]
            weights = 1.0 / np.maximum(nearest_d, 1e-6)
            preds.append((nearest_y * weights).sum(axis=1) / weights.sum(axis=1))
        return np.maximum(0.0, np.concatenate(preds))


def ensure_dirs() -> None:
    for path in [DATA_DIR, RESULTS_DIR, FIG_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def read_source() -> pd.DataFrame:
    df = pd.read_csv(SOURCE_DATA)
    df["date"] = pd.to_datetime(df["date"])
    df = df.sort_values("date").reset_index(drop=True)
    return df


def add_calendar_features(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["date"] = pd.to_datetime(out["date"])
    out["month_sin"] = np.sin(2 * np.pi * out["month"] / 12)
    out["month_cos"] = np.cos(2 * np.pi * out["month"] / 12)
    out["day_sin"] = np.sin(2 * np.pi * out["day_of_year"] / 366)
    out["day_cos"] = np.cos(2 * np.pi * out["day_of_year"] / 366)
    wd_rad = np.deg2rad(pd.to_numeric(out["WindDirection"], errors="coerce"))
    out["winddir_sin"] = np.sin(wd_rad)
    out["winddir_cos"] = np.cos(wd_rad)
    return out


def add_daily_lag_features(df: pd.DataFrame) -> pd.DataFrame:
    out = add_calendar_features(df)
    lag_vars = [
        "PM2.5",
        "AQI",
        "Temperature",
        "Humidity",
        "WindSpeed",
        "Precipitation",
        "IsRain",
        "log1p_Precipitation",
        "TempRange",
        "T_max",
        "T_min",
        "winddir_sin",
        "winddir_cos",
        "CO",
        "NO2",
        "SO2",
        "O3",
        "PM10",
    ]
    for name in lag_vars:
        for lag in range(1, 8):
            out[f"{name}_lag{lag}"] = out[name].shift(lag)

    rolling_vars = ["PM2.5", "CO", "NO2", "SO2", "O3", "PM10", "WindSpeed", "Humidity", "Precipitation"]
    for name in rolling_vars:
        shifted = out[name].shift(1)
        out[f"{name}_roll3_mean_lag1"] = shifted.rolling(3).mean()
        out[f"{name}_roll7_mean_lag1"] = shifted.rolling(7).mean()
        out[f"{name}_roll14_mean_lag1"] = shifted.rolling(14).mean()
        out[f"{name}_roll7_std_lag1"] = shifted.rolling(7).std()
        out[f"{name}_roll14_std_lag1"] = shifted.rolling(14).std()
        out[f"{name}_roll7_slope_lag1"] = shifted.rolling(7).apply(linear_slope, raw=True)

    pm_shift = out["PM2.5"].shift(1)
    out["PM2.5_7d_slope_lag1"] = pm_shift.rolling(7).apply(linear_slope, raw=True)
    out["PM2.5_3d_delta_lag1"] = out["PM2.5"].shift(1) - out["PM2.5"].shift(4)
    out["PM2.5_7d_delta_lag1"] = out["PM2.5"].shift(1) - out["PM2.5"].shift(8)
    return out


def linear_slope(values: np.ndarray) -> float:
    values = np.asarray(values, dtype=float)
    if len(values) < 2 or np.any(~np.isfinite(values)):
        return np.nan
    x = np.arange(len(values), dtype=float)
    x = x - x.mean()
    y = values - values.mean()
    denom = float(np.sum(x * x))
    return float(np.sum(x * y) / denom) if denom else 0.0


def build_daily_dataset(df: pd.DataFrame) -> tuple[pd.DataFrame, list[str], list[str]]:
    data = add_daily_lag_features(df)

    strict_weather_features = [
        "Temperature_lag1",
        "Humidity_lag1",
        "WindSpeed_lag1",
        "Precipitation_lag1",
        "IsRain_lag1",
        "log1p_Precipitation_lag1",
        "TempRange_lag1",
        "T_max_lag1",
        "T_min_lag1",
        "winddir_sin_lag1",
        "winddir_cos_lag1",
        "month_sin",
        "month_cos",
        "day_sin",
        "day_cos",
        "season_Spring",
        "season_Summer",
        "season_Autumn",
        "season_Winter",
    ]

    all_lag_vars = [
        "PM2.5",
        "AQI",
        "Temperature",
        "Humidity",
        "WindSpeed",
        "Precipitation",
        "IsRain",
        "log1p_Precipitation",
        "TempRange",
        "T_max",
        "T_min",
        "winddir_sin",
        "winddir_cos",
        "CO",
        "NO2",
        "SO2",
        "O3",
        "PM10",
    ]
    lagged_all_features = [f"{name}_lag{lag}" for name in all_lag_vars for lag in range(1, 8)]
    rolling_features = [
        col
        for col in data.columns
        if any(
            col.startswith(f"{name}_roll")
            for name in ["PM2.5", "CO", "NO2", "SO2", "O3", "PM10", "WindSpeed", "Humidity", "Precipitation"]
        )
    ]
    enhanced_features = sorted(
        set(
            strict_weather_features
            + lagged_all_features
            + rolling_features
            + [
                "PM2.5_7d_slope_lag1",
                "PM2.5_3d_delta_lag1",
                "PM2.5_7d_delta_lag1",
                "month_sin",
                "month_cos",
                "day_sin",
                "day_cos",
            ]
        )
    )

    needed = ["date", "PM2.5"] + sorted(set(enhanced_features))
    data = data[needed].replace([np.inf, -np.inf], np.nan).dropna().reset_index(drop=True)
    return data, strict_weather_features, enhanced_features


def compact_enhanced_features(strict_weather_features: list[str]) -> list[str]:
    return strict_weather_features + [
        "PM2.5_lag1",
        "PM2.5_lag2",
        "PM2.5_lag3",
        "PM2.5_roll3_mean_lag1",
        "PM2.5_roll7_mean_lag1",
        "PM2.5_roll7_std_lag1",
        "PM2.5_7d_slope_lag1",
        "PM2.5_3d_delta_lag1",
        "AQI_lag1",
        "CO_lag1",
        "NO2_lag1",
        "SO2_lag1",
        "O3_lag1",
        "PM10_lag1",
        "CO_roll3_mean_lag1",
        "NO2_roll3_mean_lag1",
        "O3_roll3_mean_lag1",
        "PM10_roll3_mean_lag1",
        "WindSpeed_roll7_mean_lag1",
        "Humidity_roll7_mean_lag1",
        "Precipitation_roll7_mean_lag1",
    ]


def split_train_valid_test(frame: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    train_valid = frame[frame["date"] < TEST_START].copy()
    test = frame[frame["date"] >= TEST_START].copy()
    valid_size = max(90, int(len(train_valid) * 0.2))
    train = train_valid.iloc[:-valid_size].copy()
    valid = train_valid.iloc[-valid_size:].copy()
    return train, valid, test


def fit_ridge(train: pd.DataFrame, features: list[str], target: str, alpha: float) -> RidgeModel:
    x = train[features].astype(float).to_numpy()
    y = train[target].astype(float).to_numpy()
    means = x.mean(axis=0)
    scales = x.std(axis=0, ddof=0)
    scales[scales == 0] = 1.0
    x_std = (x - means) / scales
    xb = np.column_stack([np.ones(len(x_std)), x_std])
    penalty = np.eye(xb.shape[1]) * alpha
    penalty[0, 0] = 0.0
    beta = np.linalg.pinv(xb.T @ xb + penalty) @ xb.T @ y
    return RidgeModel(features, means, scales, beta, alpha)


def standardize_train(train: pd.DataFrame, features: list[str]) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    x = train[features].astype(float).to_numpy()
    means = x.mean(axis=0)
    scales = x.std(axis=0, ddof=0)
    scales[scales == 0] = 1.0
    return (x - means) / scales, means, scales


def build_tree(
    x: np.ndarray,
    y: np.ndarray,
    rng: np.random.Generator,
    depth: int,
    max_depth: int,
    min_leaf: int,
    max_features: int,
    n_thresholds: int,
) -> TreeNode:
    node_value = float(np.mean(y))
    if depth >= max_depth or len(y) < 2 * min_leaf or np.var(y) < 1e-10:
        return TreeNode(value=node_value)

    feature_count = x.shape[1]
    chosen_features = rng.choice(feature_count, size=min(max_features, feature_count), replace=False)
    best_feature = None
    best_threshold = None
    best_loss = float("inf")
    best_left = None

    for feature_idx in chosen_features:
        col = x[:, feature_idx]
        if np.nanmax(col) - np.nanmin(col) < 1e-10:
            continue
        thresholds = np.unique(np.quantile(col, np.linspace(0.08, 0.92, n_thresholds)))
        for threshold in thresholds:
            left = col <= threshold
            left_count = int(left.sum())
            right_count = len(y) - left_count
            if left_count < min_leaf or right_count < min_leaf:
                continue
            y_left = y[left]
            y_right = y[~left]
            loss = float(np.sum((y_left - y_left.mean()) ** 2) + np.sum((y_right - y_right.mean()) ** 2))
            if loss < best_loss:
                best_loss = loss
                best_feature = int(feature_idx)
                best_threshold = float(threshold)
                best_left = left

    if best_feature is None or best_left is None:
        return TreeNode(value=node_value)

    return TreeNode(
        value=node_value,
        feature_idx=best_feature,
        threshold=best_threshold,
        left=build_tree(x[best_left], y[best_left], rng, depth + 1, max_depth, min_leaf, max_features, n_thresholds),
        right=build_tree(x[~best_left], y[~best_left], rng, depth + 1, max_depth, min_leaf, max_features, n_thresholds),
    )


def predict_tree(tree: TreeNode, x: np.ndarray) -> np.ndarray:
    preds = np.empty(x.shape[0], dtype=float)
    for i, row in enumerate(x):
        node = tree
        while node.feature_idx is not None and node.threshold is not None and node.left is not None and node.right is not None:
            node = node.left if row[node.feature_idx] <= node.threshold else node.right
        preds[i] = node.value
    return preds


def resolve_max_features(value: str | float | int, n_features: int) -> int:
    if value == "sqrt":
        return max(1, int(math.sqrt(n_features)))
    if isinstance(value, float):
        return max(1, int(round(n_features * value)))
    return max(1, int(value))


def fit_random_forest(
    train: pd.DataFrame,
    features: list[str],
    target: str,
    params: dict[str, float | int | str],
    seed_offset: int = 0,
) -> RandomForestModel:
    x_std, means, scales = standardize_train(train, features)
    y = train[target].astype(float).to_numpy()
    rng = np.random.default_rng(RNG_SEED + seed_offset)
    n_trees = int(params["n_trees"])
    max_depth = int(params["max_depth"])
    min_leaf = int(params["min_leaf"])
    n_thresholds = int(params["n_thresholds"])
    max_features = resolve_max_features(params["max_features"], len(features))
    trees = []
    for _ in range(n_trees):
        sample_idx = rng.integers(0, len(y), size=len(y))
        trees.append(build_tree(x_std[sample_idx], y[sample_idx], rng, 0, max_depth, min_leaf, max_features, n_thresholds))
    return RandomForestModel(features, means, scales, trees, params)


def fit_forest_with_validation(
    train: pd.DataFrame,
    valid: pd.DataFrame,
    train_valid: pd.DataFrame,
    features: list[str],
    target: str,
) -> tuple[RandomForestModel, pd.DataFrame, np.ndarray]:
    rows = []
    best_idx = 0
    best_rmse = float("inf")
    best_valid_pred = None
    for idx, params in enumerate(FOREST_CONFIGS):
        model = fit_random_forest(train, features, target, params, seed_offset=1000 + idx * 100)
        pred = model.predict(valid)
        row = metric_row("validation", f"forest_config_{idx + 1}", valid[target], pred)
        row.update(params)
        rows.append(row)
        if row["RMSE"] < best_rmse:
            best_rmse = float(row["RMSE"])
            best_idx = idx
            best_valid_pred = pred
    assert best_valid_pred is not None
    final_model = fit_random_forest(train_valid, features, target, FOREST_CONFIGS[best_idx], seed_offset=2000 + best_idx * 100)
    return final_model, pd.DataFrame(rows), best_valid_pred


def fit_knn(train: pd.DataFrame, features: list[str], target: str, k: int) -> KNNAnalogModel:
    x_std, means, scales = standardize_train(train, features)
    y = train[target].astype(float).to_numpy()
    return KNNAnalogModel(features, means, scales, x_std, y, k)


def fit_knn_with_validation(
    train: pd.DataFrame,
    valid: pd.DataFrame,
    train_valid: pd.DataFrame,
    features: list[str],
    target: str,
) -> tuple[KNNAnalogModel, pd.DataFrame, np.ndarray]:
    rows = []
    best_k = KNN_K_VALUES[0]
    best_rmse = float("inf")
    best_valid_pred = None
    for k in KNN_K_VALUES:
        model = fit_knn(train, features, target, k)
        pred = model.predict(valid)
        row = metric_row("validation", f"knn_k_{k}", valid[target], pred)
        row["k"] = k
        rows.append(row)
        if row["RMSE"] < best_rmse:
            best_rmse = float(row["RMSE"])
            best_k = k
            best_valid_pred = pred
    assert best_valid_pred is not None
    final_model = fit_knn(train_valid, features, target, best_k)
    return final_model, pd.DataFrame(rows), best_valid_pred


def select_ensemble_weights(y_true: Iterable[float], prediction_map: dict[str, np.ndarray]) -> tuple[dict[str, float], pd.DataFrame]:
    names = list(prediction_map)
    y = np.asarray(list(y_true), dtype=float)
    grid = np.arange(0, 1 + 1e-9, ENSEMBLE_WEIGHT_STEP)
    best_weights = {name: 0.0 for name in names}
    best_rmse = float("inf")
    rows = []
    for raw_weights in product(grid, repeat=len(names)):
        total = sum(raw_weights)
        if abs(total - 1.0) > 1e-9:
            continue
        pred = np.zeros_like(y, dtype=float)
        for name, weight in zip(names, raw_weights):
            pred += weight * prediction_map[name]
        rmse = math.sqrt(float(np.mean((y - pred) ** 2)))
        mae = float(np.mean(np.abs(y - pred)))
        rows.append({"RMSE": rmse, "MAE": mae, **{f"w_{name}": weight for name, weight in zip(names, raw_weights)}})
        if rmse < best_rmse:
            best_rmse = rmse
            best_weights = {name: float(weight) for name, weight in zip(names, raw_weights)}
    return best_weights, pd.DataFrame(rows).sort_values("RMSE").head(20)


def ensemble_predict(prediction_map: dict[str, np.ndarray], weights: dict[str, float]) -> np.ndarray:
    first = next(iter(prediction_map.values()))
    pred = np.zeros_like(first, dtype=float)
    for name, values in prediction_map.items():
        pred += weights.get(name, 0.0) * values
    return np.maximum(0.0, pred)


def fit_ridge_with_validation(
    train: pd.DataFrame,
    valid: pd.DataFrame,
    train_valid: pd.DataFrame,
    features: list[str],
    target: str,
) -> tuple[RidgeModel, pd.DataFrame]:
    rows = []
    best_alpha = ALPHAS[0]
    best_rmse = float("inf")
    for alpha in ALPHAS:
        model = fit_ridge(train, features, target, alpha)
        pred = model.predict(valid)
        row = metric_row("validation", f"ridge_alpha_{alpha:g}", valid[target].to_numpy(), pred)
        row["alpha"] = alpha
        rows.append(row)
        if row["RMSE"] < best_rmse:
            best_rmse = row["RMSE"]
            best_alpha = alpha
    final_model = fit_ridge(train_valid, features, target, best_alpha)
    return final_model, pd.DataFrame(rows)


def metric_row(split: str, model: str, y_true: Iterable[float], y_pred: Iterable[float]) -> dict[str, float | str]:
    y = np.asarray(list(y_true), dtype=float)
    pred = np.asarray(list(y_pred), dtype=float)
    residual = y - pred
    sse = float(np.sum(residual**2))
    sst = float(np.sum((y - y.mean()) ** 2))
    return {
        "split": split,
        "model": model,
        "n": int(len(y)),
        "RMSE": math.sqrt(float(np.mean(residual**2))),
        "MAE": float(np.mean(np.abs(residual))),
        "Bias": float(np.mean(pred - y)),
        "R2": 1.0 - sse / sst if sst else np.nan,
    }


def normal_cdf(value: float) -> float:
    return 0.5 * (1.0 + math.erf(value / math.sqrt(2.0)))


def normal_two_sided_p(value: float) -> float:
    return math.erfc(abs(value) / math.sqrt(2.0))


def wilson_interval(successes: int, n: int, z: float = 1.96) -> tuple[float, float]:
    if n == 0:
        return np.nan, np.nan
    phat = successes / n
    denom = 1 + z**2 / n
    centre = (phat + z**2 / (2 * n)) / denom
    half = z * math.sqrt((phat * (1 - phat) + z**2 / (4 * n)) / n) / denom
    return max(0.0, centre - half), min(1.0, centre + half)


def bootstrap_metric_ci(
    y_true: Iterable[float],
    y_pred: Iterable[float],
    model_name: str,
    n_boot: int = 1000,
    seed_offset: int = 0,
) -> dict[str, float | str]:
    y = np.asarray(list(y_true), dtype=float)
    pred = np.asarray(list(y_pred), dtype=float)
    n = len(y)
    rng = np.random.default_rng(RNG_SEED + seed_offset)
    rows = np.empty((n_boot, 3), dtype=float)
    for i in range(n_boot):
        idx = rng.integers(0, n, size=n)
        metrics = metric_row("bootstrap", model_name, y[idx], pred[idx])
        rows[i, 0] = float(metrics["RMSE"])
        rows[i, 1] = float(metrics["MAE"])
        rows[i, 2] = float(metrics["R2"])
    base = metric_row("test", model_name, y, pred)
    return {
        "model": model_name,
        "n": n,
        "RMSE": base["RMSE"],
        "RMSE_ci_low": float(np.quantile(rows[:, 0], 0.025)),
        "RMSE_ci_high": float(np.quantile(rows[:, 0], 0.975)),
        "MAE": base["MAE"],
        "MAE_ci_low": float(np.quantile(rows[:, 1], 0.025)),
        "MAE_ci_high": float(np.quantile(rows[:, 1], 0.975)),
        "R2": base["R2"],
        "R2_ci_low": float(np.quantile(rows[:, 2], 0.025)),
        "R2_ci_high": float(np.quantile(rows[:, 2], 0.975)),
    }


def paired_error_test(
    y_true: Iterable[float],
    model_pred: Iterable[float],
    reference_pred: Iterable[float],
    model_name: str,
    reference_name: str,
) -> dict[str, float | str]:
    y = np.asarray(list(y_true), dtype=float)
    model = np.asarray(list(model_pred), dtype=float)
    reference = np.asarray(list(reference_pred), dtype=float)
    model_sq = (y - model) ** 2
    ref_sq = (y - reference) ** 2
    diff = ref_sq - model_sq
    mean_diff = float(diff.mean())
    se = float(diff.std(ddof=1) / math.sqrt(len(diff))) if len(diff) > 1 else np.nan
    z = mean_diff / se if se and np.isfinite(se) else np.nan
    p_value = 1.0 - normal_cdf(z) if np.isfinite(z) else np.nan
    model_rmse = math.sqrt(float(model_sq.mean()))
    reference_rmse = math.sqrt(float(ref_sq.mean()))
    conclusion = "显著更准" if np.isfinite(p_value) and p_value < 0.05 and mean_diff > 0 else "未显著更准"
    return {
        "test_type": "paired_squared_error",
        "model": model_name,
        "reference": reference_name,
        "n": len(y),
        "model_RMSE": model_rmse,
        "reference_RMSE": reference_rmse,
        "RMSE_reduction_percent": (reference_rmse - model_rmse) / reference_rmse * 100 if reference_rmse else np.nan,
        "mean_squared_error_reduction": mean_diff,
        "z_statistic": z,
        "p_value_one_sided": p_value,
        "conclusion": conclusion,
    }


def residual_bias_test(y_true: Iterable[float], y_pred: Iterable[float], model_name: str) -> dict[str, float | str]:
    y = np.asarray(list(y_true), dtype=float)
    pred = np.asarray(list(y_pred), dtype=float)
    bias = pred - y
    n = len(bias)
    mean_bias = float(bias.mean())
    se = float(bias.std(ddof=1) / math.sqrt(n)) if n > 1 else np.nan
    z = mean_bias / se if se and np.isfinite(se) else np.nan
    p_value = normal_two_sided_p(z) if np.isfinite(z) else np.nan
    low = mean_bias - 1.96 * se if np.isfinite(se) else np.nan
    high = mean_bias + 1.96 * se if np.isfinite(se) else np.nan
    conclusion = "存在系统性偏差" if np.isfinite(p_value) and p_value < 0.05 else "无显著系统性偏差"
    return {
        "test_type": "residual_bias",
        "model": model_name,
        "reference": "",
        "n": n,
        "bias_mean": mean_bias,
        "bias_ci_low": low,
        "bias_ci_high": high,
        "z_statistic": z,
        "p_value_two_sided": p_value,
        "conclusion": conclusion,
    }


def conformal_interval_validation(
    y_calib: Iterable[float],
    pred_calib: Iterable[float],
    y_test: Iterable[float],
    pred_test: Iterable[float],
    model_name: str,
    target_coverage: float = 0.90,
) -> dict[str, float | str]:
    calib_abs = np.sort(np.abs(np.asarray(list(y_calib), dtype=float) - np.asarray(list(pred_calib), dtype=float)))
    test_y = np.asarray(list(y_test), dtype=float)
    test_pred = np.asarray(list(pred_test), dtype=float)
    q_idx = int(math.ceil((len(calib_abs) + 1) * target_coverage) - 1)
    q_idx = min(max(q_idx, 0), len(calib_abs) - 1)
    half_width = float(calib_abs[q_idx])
    lower = np.maximum(0.0, test_pred - half_width)
    upper = test_pred + half_width
    coverage = float(np.mean((test_y >= lower) & (test_y <= upper)))
    return {
        "model": model_name,
        "target_coverage": target_coverage,
        "empirical_coverage": coverage,
        "interval_half_width": half_width,
        "average_interval_width": float(np.mean(upper - lower)),
        "covered_days": int(np.sum((test_y >= lower) & (test_y <= upper))),
        "n": len(test_y),
    }


def classify_levels(values: Iterable[float], q33: float, q67: float) -> np.ndarray:
    arr = np.asarray(list(values), dtype=float)
    labels = np.full(arr.shape, "Medium", dtype=object)
    labels[arr < q33] = "Low"
    labels[arr >= q67] = "High"
    return labels


def classification_metrics(y_true: Iterable[float], y_pred: Iterable[float], q33: float, q67: float, model_name: str) -> dict[str, float | str]:
    actual = classify_levels(y_true, q33, q67)
    pred = classify_levels(y_pred, q33, q67)
    labels = ["Low", "Medium", "High"]
    accuracy = float(np.mean(actual == pred))

    f1_values = []
    support_values = []
    precision_by_label = {}
    recall_by_label = {}
    for label in labels:
        tp = int(np.sum((actual == label) & (pred == label)))
        fp = int(np.sum((actual != label) & (pred == label)))
        fn = int(np.sum((actual == label) & (pred != label)))
        support = int(np.sum(actual == label))
        precision = tp / (tp + fp) if (tp + fp) else 0.0
        recall = tp / (tp + fn) if (tp + fn) else 0.0
        f1 = 2 * precision * recall / (precision + recall) if (precision + recall) else 0.0
        f1_values.append(f1)
        support_values.append(support)
        precision_by_label[label] = precision
        recall_by_label[label] = recall

    weighted_f1 = float(np.average(f1_values, weights=support_values)) if sum(support_values) else 0.0
    return {
        "model": model_name,
        "level_threshold_low_medium": q33,
        "level_threshold_medium_high": q67,
        "accuracy_3class": accuracy,
        "weighted_f1_3class": weighted_f1,
        "high_precision": precision_by_label["High"],
        "high_recall": recall_by_label["High"],
        "high_f1": f1_values[2],
    }


def build_confusion_matrix(y_true: Iterable[float], y_pred: Iterable[float], q33: float, q67: float, model_name: str) -> pd.DataFrame:
    actual = classify_levels(y_true, q33, q67)
    pred = classify_levels(y_pred, q33, q67)
    labels = ["Low", "Medium", "High"]
    rows = []
    for a in labels:
        for p in labels:
            rows.append({"model": model_name, "actual_level": a, "predicted_level": p, "count": int(np.sum((actual == a) & (pred == p)))})
    return pd.DataFrame(rows)


def classification_accuracy_ci(y_true: Iterable[float], y_pred: Iterable[float], q33: float, q67: float, model_name: str) -> dict[str, float | str]:
    actual = classify_levels(y_true, q33, q67)
    pred = classify_levels(y_pred, q33, q67)
    successes = int(np.sum(actual == pred))
    low, high = wilson_interval(successes, len(actual))
    return {
        "model": model_name,
        "n": len(actual),
        "accuracy_3class": successes / len(actual) if len(actual) else np.nan,
        "accuracy_ci_low": low,
        "accuracy_ci_high": high,
        "correct_days": successes,
    }


def add_trend_anchor_features(df: pd.DataFrame) -> pd.DataFrame:
    base = add_calendar_features(df)
    records = []
    value_cols = ["PM2.5", "Temperature", "Humidity", "WindSpeed", "Precipitation", "CO", "NO2", "SO2", "O3", "PM10"]
    for i in range(6, len(base) - 7):
        row = {
            "anchor_index": i,
            "anchor_date": base.loc[i, "date"],
            "anchor_PM2.5": base.loc[i, "PM2.5"],
        }
        history = base.iloc[i - 6 : i + 1].copy()
        for lag in range(0, 7):
            hist_row = base.loc[i - lag]
            for col in value_cols:
                row[f"{col}_lag{lag}"] = hist_row[col]
        pm_values = history["PM2.5"].to_numpy()
        row["PM2.5_window7_mean"] = float(np.mean(pm_values))
        row["PM2.5_window7_std"] = float(np.std(pm_values, ddof=0))
        row["PM2.5_window7_min"] = float(np.min(pm_values))
        row["PM2.5_window7_max"] = float(np.max(pm_values))
        row["PM2.5_window7_slope"] = linear_slope(pm_values)
        for col in ["Temperature", "Humidity", "WindSpeed", "Precipitation", "CO", "NO2", "O3", "PM10"]:
            row[f"{col}_window7_mean"] = float(history[col].mean())
            row[f"{col}_window7_slope"] = linear_slope(history[col].to_numpy())
        for horizon in range(1, 8):
            target_row = base.loc[i + horizon]
            row[f"target_h{horizon}"] = target_row["PM2.5"]
            row[f"target_date_h{horizon}"] = target_row["date"]
            day_of_year = int(target_row["day_of_year"])
            month = int(target_row["month"])
            row[f"target_day_sin_h{horizon}"] = math.sin(2 * math.pi * day_of_year / 366)
            row[f"target_day_cos_h{horizon}"] = math.cos(2 * math.pi * day_of_year / 366)
            row[f"target_month_sin_h{horizon}"] = math.sin(2 * math.pi * month / 12)
            row[f"target_month_cos_h{horizon}"] = math.cos(2 * math.pi * month / 12)
        records.append(row)
    return pd.DataFrame(records).replace([np.inf, -np.inf], np.nan).dropna().reset_index(drop=True)


def trend_base_features(anchor_frame: pd.DataFrame) -> list[str]:
    excluded_prefixes = ("target_h", "target_date_h", "target_day_", "target_month_")
    excluded = {"anchor_index", "anchor_date", "anchor_PM2.5"}
    return [
        col
        for col in anchor_frame.columns
        if col not in excluded and not col.startswith(excluded_prefixes)
    ]


def fit_trend_models(anchor_frame: pd.DataFrame, q33: float, q67: float) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    base_features = trend_base_features(anchor_frame)
    predictions = []
    metrics = []
    val_rows = []
    coefficient_frames = []

    for horizon in range(1, 8):
        target = f"target_h{horizon}"
        target_date = f"target_date_h{horizon}"
        frame = anchor_frame.copy()
        frame["target_date"] = pd.to_datetime(frame[target_date])
        features = base_features + [
            f"target_day_sin_h{horizon}",
            f"target_day_cos_h{horizon}",
            f"target_month_sin_h{horizon}",
            f"target_month_cos_h{horizon}",
        ]
        model_frame = frame[["anchor_date", "anchor_PM2.5", "target_date", target] + features].rename(columns={target: "PM2.5"}).dropna().copy()
        train_valid = model_frame[model_frame["target_date"] < TEST_START].copy()
        test = model_frame[model_frame["target_date"] >= TEST_START].copy()
        valid_size = max(70, int(len(train_valid) * 0.2))
        train = train_valid.iloc[:-valid_size].copy()
        valid = train_valid.iloc[-valid_size:].copy()
        model, val = fit_ridge_with_validation(train, valid, train_valid, features, "PM2.5")
        val["horizon"] = horizon
        val_rows.append(val)
        coefficient_frames.append(model.coefficient_frame(f"trend_h{horizon}", top_n=15))
        valid_ridge_model = fit_ridge(train, features, "PM2.5", model.alpha)
        ridge_valid_pred = valid_ridge_model.predict(valid)
        knn_model, knn_val, knn_valid_pred = fit_knn_with_validation(train, valid, train_valid, features, "PM2.5")
        knn_val["horizon"] = horizon
        val_rows.append(knn_val.assign(model_family=f"trend_knn_h{horizon}"))
        valid_persistence = np.maximum(0.0, valid["anchor_PM2.5"].to_numpy())
        trend_weights, trend_grid = select_ensemble_weights(
            valid["PM2.5"].to_numpy(),
            {"ridge": ridge_valid_pred, "analog": knn_valid_pred, "persistence": valid_persistence},
        )

        ridge_pred = model.predict(test)
        knn_pred = knn_model.predict(test)
        baseline = np.maximum(0.0, test["anchor_PM2.5"].to_numpy())
        pred = ensemble_predict({"ridge": ridge_pred, "analog": knn_pred, "persistence": baseline}, trend_weights)
        metrics.append(metric_row("test", f"Trend optimized ensemble H{horizon}", test["PM2.5"].to_numpy(), pred) | {"horizon": horizon, "alpha": np.nan, "k": knn_model.k})
        metrics.append(metric_row("test", f"Trend direct ridge H{horizon}", test["PM2.5"].to_numpy(), ridge_pred) | {"horizon": horizon, "alpha": model.alpha})
        metrics.append(metric_row("test", f"Trend analog KNN H{horizon}", test["PM2.5"].to_numpy(), knn_pred) | {"horizon": horizon, "alpha": np.nan, "k": knn_model.k})
        metrics.append(metric_row("test", f"Persistence baseline H{horizon}", test["PM2.5"].to_numpy(), baseline) | {"horizon": horizon, "alpha": np.nan})
        actual_direction = np.sign(test["PM2.5"].to_numpy() - test["anchor_PM2.5"].to_numpy())
        pred_direction = np.sign(pred - test["anchor_PM2.5"].to_numpy())
        direction_accuracy = float(np.mean(actual_direction == pred_direction))
        ridge_direction_accuracy = float(np.mean(actual_direction == np.sign(ridge_pred - test["anchor_PM2.5"].to_numpy())))
        knn_direction_accuracy = float(np.mean(actual_direction == np.sign(knn_pred - test["anchor_PM2.5"].to_numpy())))
        baseline_direction_accuracy = float(np.mean(actual_direction == np.sign(baseline - test["anchor_PM2.5"].to_numpy())))
        metrics[-4]["direction_accuracy_vs_anchor"] = direction_accuracy
        metrics[-3]["direction_accuracy_vs_anchor"] = ridge_direction_accuracy
        metrics[-2]["direction_accuracy_vs_anchor"] = knn_direction_accuracy
        metrics[-1]["direction_accuracy_vs_anchor"] = baseline_direction_accuracy
        for name, weight in trend_weights.items():
            metrics[-4][f"w_{name}"] = weight

        out = test[["anchor_date", "target_date", "anchor_PM2.5", "PM2.5"]].copy()
        out["horizon"] = horizon
        out["predicted_PM2.5"] = pred
        out["ridge_predicted_PM2.5"] = ridge_pred
        out["analog_predicted_PM2.5"] = knn_pred
        out["baseline_PM2.5"] = baseline
        out["predicted_level"] = classify_levels(pred, q33, q67)
        out["actual_level"] = classify_levels(out["PM2.5"].to_numpy(), q33, q67)
        predictions.append(out)

    return (
        pd.concat(predictions, ignore_index=True),
        pd.DataFrame(metrics),
        pd.concat(coefficient_frames, ignore_index=True),
    )


def write_daily_outputs(
    daily_data: pd.DataFrame,
    strict_features: list[str],
    enhanced_features: list[str],
) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, dict[str, pd.DataFrame], float, float]:
    train, valid, test = split_train_valid_test(daily_data)
    train_valid = daily_data[daily_data["date"] < TEST_START].copy()
    q33 = float(train_valid["PM2.5"].quantile(1 / 3))
    q67 = float(train_valid["PM2.5"].quantile(2 / 3))
    compact_features = compact_enhanced_features(strict_features)

    weather_model, weather_val = fit_ridge_with_validation(train, valid, train_valid, strict_features, "PM2.5")
    compact_model, compact_val = fit_ridge_with_validation(train, valid, train_valid, compact_features, "PM2.5")
    enhanced_model, enhanced_val = fit_ridge_with_validation(train, valid, train_valid, enhanced_features, "PM2.5")
    forest_model, forest_val, forest_valid_pred = fit_forest_with_validation(train, valid, train_valid, enhanced_features, "PM2.5")
    knn_model, knn_val, knn_valid_pred = fit_knn_with_validation(train, valid, train_valid, enhanced_features, "PM2.5")

    weather_valid_model = fit_ridge(train, strict_features, "PM2.5", weather_model.alpha)
    weather_valid_pred = weather_valid_model.predict(valid)
    compact_valid_model = fit_ridge(train, compact_features, "PM2.5", compact_model.alpha)
    compact_valid_pred = compact_valid_model.predict(valid)
    enhanced_valid_model = fit_ridge(train, enhanced_features, "PM2.5", enhanced_model.alpha)
    enhanced_valid_pred = enhanced_valid_model.predict(valid)
    valid_persistence_pred = np.maximum(0.0, valid["PM2.5_lag1"].to_numpy())
    ensemble_weights, ensemble_validation_grid = select_ensemble_weights(
        valid["PM2.5"].to_numpy(),
        {
            "compact": compact_valid_pred,
            "all_lag": enhanced_valid_pred,
            "forest": forest_valid_pred,
            "analog": knn_valid_pred,
            "persistence": valid_persistence_pred,
        },
    )

    weather_pred = weather_model.predict(test)
    compact_pred = compact_model.predict(test)
    enhanced_pred = enhanced_model.predict(test)
    forest_pred = forest_model.predict(test)
    knn_pred = knn_model.predict(test)
    persistence_pred = np.maximum(0.0, test["PM2.5_lag1"].to_numpy())
    mean_pred = np.full(len(test), train_valid["PM2.5"].mean())
    valid_mean_pred = np.full(len(valid), train["PM2.5"].mean())
    validation_prediction_map = {
        "Weather-only previous-day ridge": weather_valid_pred,
        "Compact enhanced ridge": compact_valid_pred,
        "All-lag enhanced ridge": enhanced_valid_pred,
        "Enhanced lag random forest": forest_valid_pred,
        "Historical analog KNN": knn_valid_pred,
        "Persistence baseline": valid_persistence_pred,
        "Train-mean baseline": valid_mean_pred,
    }
    validation_prediction_map["Validation-weighted ensemble"] = ensemble_predict(
        {
            "compact": compact_valid_pred,
            "all_lag": enhanced_valid_pred,
            "forest": forest_valid_pred,
            "analog": knn_valid_pred,
            "persistence": valid_persistence_pred,
        },
        ensemble_weights,
    )
    ensemble_pred = ensemble_predict(
        {
            "compact": compact_pred,
            "all_lag": enhanced_pred,
            "forest": forest_pred,
            "analog": knn_pred,
            "persistence": persistence_pred,
        },
        ensemble_weights,
    )
    test_prediction_map = {
        "Weather-only previous-day ridge": weather_pred,
        "Compact enhanced ridge": compact_pred,
        "All-lag enhanced ridge": enhanced_pred,
        "Enhanced lag random forest": forest_pred,
        "Historical analog KNN": knn_pred,
        "Validation-weighted ensemble": ensemble_pred,
        "Persistence baseline": persistence_pred,
        "Train-mean baseline": mean_pred,
    }

    metrics = pd.DataFrame(
        [
            metric_row("test", "Weather-only previous-day ridge", test["PM2.5"], weather_pred) | {"alpha": weather_model.alpha},
            metric_row("test", "Compact enhanced ridge", test["PM2.5"], compact_pred) | {"alpha": compact_model.alpha},
            metric_row("test", "All-lag enhanced ridge", test["PM2.5"], enhanced_pred) | {"alpha": enhanced_model.alpha},
            metric_row("test", "Enhanced lag random forest", test["PM2.5"], forest_pred) | {"alpha": np.nan},
            metric_row("test", "Historical analog KNN", test["PM2.5"], knn_pred) | {"alpha": np.nan, "k": knn_model.k},
            metric_row("test", "Validation-weighted ensemble", test["PM2.5"], ensemble_pred) | {"alpha": np.nan},
            metric_row("test", "Persistence baseline", test["PM2.5"], persistence_pred) | {"alpha": np.nan},
            metric_row("test", "Train-mean baseline", test["PM2.5"], mean_pred) | {"alpha": np.nan},
        ]
    )
    for name, weight in ensemble_weights.items():
        metrics.loc[metrics["model"] == "Validation-weighted ensemble", f"w_{name}"] = weight

    class_rows = pd.DataFrame(
        [
            classification_metrics(test["PM2.5"], weather_pred, q33, q67, "Weather-only previous-day ridge"),
            classification_metrics(test["PM2.5"], compact_pred, q33, q67, "Compact enhanced ridge"),
            classification_metrics(test["PM2.5"], enhanced_pred, q33, q67, "All-lag enhanced ridge"),
            classification_metrics(test["PM2.5"], forest_pred, q33, q67, "Enhanced lag random forest"),
            classification_metrics(test["PM2.5"], knn_pred, q33, q67, "Historical analog KNN"),
            classification_metrics(test["PM2.5"], ensemble_pred, q33, q67, "Validation-weighted ensemble"),
            classification_metrics(test["PM2.5"], persistence_pred, q33, q67, "Persistence baseline"),
        ]
    )
    confusion = pd.concat(
        [
            build_confusion_matrix(test["PM2.5"], weather_pred, q33, q67, "Weather-only previous-day ridge"),
            build_confusion_matrix(test["PM2.5"], compact_pred, q33, q67, "Compact enhanced ridge"),
            build_confusion_matrix(test["PM2.5"], enhanced_pred, q33, q67, "All-lag enhanced ridge"),
            build_confusion_matrix(test["PM2.5"], forest_pred, q33, q67, "Enhanced lag random forest"),
            build_confusion_matrix(test["PM2.5"], knn_pred, q33, q67, "Historical analog KNN"),
            build_confusion_matrix(test["PM2.5"], ensemble_pred, q33, q67, "Validation-weighted ensemble"),
            build_confusion_matrix(test["PM2.5"], persistence_pred, q33, q67, "Persistence baseline"),
        ],
        ignore_index=True,
    )

    bootstrap_ci = pd.DataFrame(
        [
            bootstrap_metric_ci(test["PM2.5"], pred, model_name, seed_offset=idx * 31)
            for idx, (model_name, pred) in enumerate(test_prediction_map.items())
        ]
    )
    paired_tests = pd.DataFrame(
        [
            paired_error_test(test["PM2.5"], test_prediction_map["Validation-weighted ensemble"], test_prediction_map["Weather-only previous-day ridge"], "Validation-weighted ensemble", "Weather-only previous-day ridge"),
            paired_error_test(test["PM2.5"], test_prediction_map["Validation-weighted ensemble"], test_prediction_map["Persistence baseline"], "Validation-weighted ensemble", "Persistence baseline"),
            paired_error_test(test["PM2.5"], test_prediction_map["Compact enhanced ridge"], test_prediction_map["Weather-only previous-day ridge"], "Compact enhanced ridge", "Weather-only previous-day ridge"),
            paired_error_test(test["PM2.5"], test_prediction_map["Compact enhanced ridge"], test_prediction_map["Persistence baseline"], "Compact enhanced ridge", "Persistence baseline"),
            paired_error_test(test["PM2.5"], test_prediction_map["All-lag enhanced ridge"], test_prediction_map["Weather-only previous-day ridge"], "All-lag enhanced ridge", "Weather-only previous-day ridge"),
            paired_error_test(test["PM2.5"], test_prediction_map["Enhanced lag random forest"], test_prediction_map["Weather-only previous-day ridge"], "Enhanced lag random forest", "Weather-only previous-day ridge"),
        ]
    )
    residual_tests = pd.DataFrame(
        [residual_bias_test(test["PM2.5"], pred, model_name) for model_name, pred in test_prediction_map.items()]
    )
    conformal_intervals = pd.DataFrame(
        [
            conformal_interval_validation(
                valid["PM2.5"],
                validation_prediction_map[model_name],
                test["PM2.5"],
                test_prediction_map[model_name],
                model_name,
            )
            for model_name in test_prediction_map
            if model_name in validation_prediction_map
        ]
    )
    classification_ci = pd.DataFrame(
        [
            classification_accuracy_ci(test["PM2.5"], pred, q33, q67, model_name)
            for model_name, pred in test_prediction_map.items()
            if model_name != "Train-mean baseline"
        ]
    )

    predictions = test[["date", "PM2.5", "PM2.5_lag1"]].copy()
    predictions["weather_only_pred_PM2.5"] = weather_pred
    predictions["compact_enhanced_pred_PM2.5"] = compact_pred
    predictions["all_lag_enhanced_pred_PM2.5"] = enhanced_pred
    predictions["random_forest_pred_PM2.5"] = forest_pred
    predictions["analog_knn_pred_PM2.5"] = knn_pred
    predictions["ensemble_pred_PM2.5"] = ensemble_pred
    predictions["persistence_pred_PM2.5"] = persistence_pred
    predictions["actual_level"] = classify_levels(predictions["PM2.5"], q33, q67)
    predictions["weather_only_pred_level"] = classify_levels(weather_pred, q33, q67)
    predictions["compact_enhanced_pred_level"] = classify_levels(compact_pred, q33, q67)
    predictions["all_lag_enhanced_pred_level"] = classify_levels(enhanced_pred, q33, q67)
    predictions["random_forest_pred_level"] = classify_levels(forest_pred, q33, q67)
    predictions["analog_knn_pred_level"] = classify_levels(knn_pred, q33, q67)
    predictions["ensemble_pred_level"] = classify_levels(ensemble_pred, q33, q67)
    predictions["persistence_pred_level"] = classify_levels(persistence_pred, q33, q67)

    coeffs = pd.concat(
        [
            weather_model.coefficient_frame("Weather-only previous-day ridge", top_n=25),
            compact_model.coefficient_frame("Compact enhanced ridge", top_n=35),
            enhanced_model.coefficient_frame("All-lag enhanced ridge", top_n=35),
        ],
        ignore_index=True,
    )

    validation = pd.concat(
        [
            weather_val.assign(model_family="Weather-only previous-day ridge"),
            compact_val.assign(model_family="Compact enhanced ridge"),
            enhanced_val.assign(model_family="All-lag enhanced ridge"),
            forest_val.assign(model_family="Enhanced lag random forest"),
            knn_val.assign(model_family="Historical analog KNN"),
        ],
        ignore_index=True,
    )

    split_summary = pd.DataFrame(
        [
            {"split": "train_core", "start_date": train["date"].min(), "end_date": train["date"].max(), "rows": len(train)},
            {"split": "validation", "start_date": valid["date"].min(), "end_date": valid["date"].max(), "rows": len(valid)},
            {"split": "test", "start_date": test["date"].min(), "end_date": test["date"].max(), "rows": len(test)},
            {"split": "train_valid_for_final_fit", "start_date": train_valid["date"].min(), "end_date": train_valid["date"].max(), "rows": len(train_valid)},
        ]
    )

    daily_data.to_csv(DATA_DIR / "prediction_feature_dataset.csv", index=False)
    predictions.to_csv(DATA_DIR / "daily_forecast_test_predictions.csv", index=False)
    metrics.to_csv(RESULTS_DIR / "daily_model_metrics.csv", index=False)
    class_rows.to_csv(RESULTS_DIR / "daily_classification_metrics.csv", index=False)
    confusion.to_csv(RESULTS_DIR / "daily_level_confusion_matrix.csv", index=False)
    coeffs.to_csv(RESULTS_DIR / "daily_model_coefficients.csv", index=False)
    validation.to_csv(RESULTS_DIR / "daily_alpha_validation.csv", index=False)
    pd.DataFrame([ensemble_weights]).to_csv(RESULTS_DIR / "daily_ensemble_weights.csv", index=False)
    ensemble_validation_grid.to_csv(RESULTS_DIR / "daily_ensemble_validation_grid_top20.csv", index=False)
    split_summary.to_csv(RESULTS_DIR / "data_split_summary.csv", index=False)
    bootstrap_ci.to_csv(RESULTS_DIR / "stat_bootstrap_metric_ci.csv", index=False)
    paired_tests.to_csv(RESULTS_DIR / "stat_paired_error_tests.csv", index=False)
    residual_tests.to_csv(RESULTS_DIR / "stat_residual_bias_tests.csv", index=False)
    conformal_intervals.to_csv(RESULTS_DIR / "stat_conformal_intervals.csv", index=False)
    classification_ci.to_csv(RESULTS_DIR / "stat_classification_wilson_ci.csv", index=False)

    stat_tables = {
        "bootstrap_ci": bootstrap_ci,
        "paired_tests": paired_tests,
        "residual_tests": residual_tests,
        "conformal_intervals": conformal_intervals,
        "classification_ci": classification_ci,
    }
    return predictions, metrics, class_rows, coeffs, stat_tables, q33, q67


def write_trend_outputs(source_df: pd.DataFrame, q33: float, q67: float) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    anchor_frame = add_trend_anchor_features(source_df)
    predictions, metrics, coeffs = fit_trend_models(anchor_frame, q33, q67)
    anchor_frame.to_csv(DATA_DIR / "trend_anchor_feature_dataset.csv", index=False)
    predictions.to_csv(DATA_DIR / "trend_forecast_test_predictions.csv", index=False)
    metrics.to_csv(RESULTS_DIR / "trend_horizon_metrics.csv", index=False)
    coeffs.to_csv(RESULTS_DIR / "trend_model_coefficients_top.csv", index=False)
    return predictions, metrics, coeffs


def load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def draw_line_chart(
    path: Path,
    title: str,
    x_labels: list[str],
    series: list[tuple[str, list[float], tuple[int, int, int]]],
    y_label: str = "PM2.5",
    width: int = 1400,
    height: int = 760,
) -> None:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = load_font(22)
    small = load_font(17)
    title_font = load_font(28, bold=True)
    margin_l, margin_r, margin_t, margin_b = 92, 42, 82, 100
    plot_w = width - margin_l - margin_r
    plot_h = height - margin_t - margin_b
    all_values = [v for _, values, _ in series for v in values if np.isfinite(v)]
    y_min = max(0, math.floor(min(all_values) / 10) * 10)
    y_max = math.ceil(max(all_values) / 10) * 10
    if y_max <= y_min:
        y_max = y_min + 10

    draw.text((margin_l, 24), title, fill=(20, 34, 54), font=title_font)
    draw.line((margin_l, margin_t, margin_l, margin_t + plot_h), fill=(90, 90, 90), width=2)
    draw.line((margin_l, margin_t + plot_h, margin_l + plot_w, margin_t + plot_h), fill=(90, 90, 90), width=2)

    for i in range(6):
        y_val = y_min + (y_max - y_min) * i / 5
        y = margin_t + plot_h - (y_val - y_min) / (y_max - y_min) * plot_h
        draw.line((margin_l, y, margin_l + plot_w, y), fill=(230, 233, 238), width=1)
        draw.text((22, y - 10), f"{y_val:.0f}", fill=(80, 80, 80), font=small)
    draw.text((20, margin_t - 32), y_label, fill=(80, 80, 80), font=small)

    n = len(x_labels)
    def xy(idx: int, val: float) -> tuple[float, float]:
        x = margin_l + idx / max(1, n - 1) * plot_w
        y = margin_t + plot_h - (val - y_min) / (y_max - y_min) * plot_h
        return x, y

    for label, values, color in series:
        points = [xy(i, float(v)) for i, v in enumerate(values)]
        if len(points) >= 2:
            draw.line(points, fill=color, width=3, joint="curve")
        for point in points[:: max(1, len(points) // 60)]:
            x, y = point
            draw.ellipse((x - 2, y - 2, x + 2, y + 2), fill=color)

    ticks = np.linspace(0, n - 1, min(7, n), dtype=int)
    for idx in ticks:
        x = margin_l + idx / max(1, n - 1) * plot_w
        draw.line((x, margin_t + plot_h, x, margin_t + plot_h + 5), fill=(90, 90, 90), width=1)
        draw.text((x - 45, margin_t + plot_h + 12), x_labels[idx], fill=(80, 80, 80), font=small)

    legend_x = margin_l
    legend_y = height - 58
    for label, _, color in series:
        draw.line((legend_x, legend_y + 11, legend_x + 34, legend_y + 11), fill=color, width=4)
        draw.text((legend_x + 42, legend_y), label, fill=(40, 40, 40), font=font)
        legend_x += 330

    img.save(path)


def draw_scatter_chart(path: Path, title: str, actual: Iterable[float], predicted: Iterable[float], width: int = 900, height: int = 760) -> None:
    y = np.asarray(list(actual), dtype=float)
    p = np.asarray(list(predicted), dtype=float)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(27, bold=True)
    font = load_font(16)
    margin = 90
    plot_w = width - 2 * margin
    plot_h = height - 2 * margin
    lim = math.ceil(max(float(y.max()), float(p.max())) / 10) * 10
    lim = max(lim, 20)
    draw.text((margin, 26), title, fill=(20, 34, 54), font=title_font)
    draw.rectangle((margin, margin, margin + plot_w, margin + plot_h), outline=(95, 95, 95), width=2)
    for i in range(6):
        val = lim * i / 5
        x = margin + val / lim * plot_w
        yy = margin + plot_h - val / lim * plot_h
        draw.line((x, margin, x, margin + plot_h), fill=(234, 236, 241))
        draw.line((margin, yy, margin + plot_w, yy), fill=(234, 236, 241))
        draw.text((x - 10, margin + plot_h + 10), f"{val:.0f}", font=font, fill=(80, 80, 80))
        draw.text((36, yy - 8), f"{val:.0f}", font=font, fill=(80, 80, 80))
    draw.line((margin, margin + plot_h, margin + plot_w, margin), fill=(40, 40, 40), width=2)
    for actual_val, pred_val in zip(y, p):
        x = margin + actual_val / lim * plot_w
        yy = margin + plot_h - pred_val / lim * plot_h
        draw.ellipse((x - 3, yy - 3, x + 3, yy + 3), fill=(39, 111, 191))
    draw.text((width // 2 - 60, height - 42), "Observed PM2.5", font=font, fill=(50, 50, 50))
    draw.text((20, 48), "Predicted", font=font, fill=(50, 50, 50))
    img.save(path)


def draw_bar_chart(path: Path, title: str, labels: list[str], values: list[float], width: int = 1050, height: int = 690) -> None:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(27, bold=True)
    font = load_font(18)
    small = load_font(15)
    margin_l, margin_r, margin_t, margin_b = 86, 40, 85, 92
    plot_w = width - margin_l - margin_r
    plot_h = height - margin_t - margin_b
    max_val = max(values) * 1.15
    draw.text((margin_l, 26), title, fill=(20, 34, 54), font=title_font)
    for i in range(6):
        val = max_val * i / 5
        y = margin_t + plot_h - val / max_val * plot_h
        draw.line((margin_l, y, margin_l + plot_w, y), fill=(232, 235, 240))
        draw.text((22, y - 9), f"{val:.1f}", fill=(80, 80, 80), font=small)
    bar_w = plot_w / len(values) * 0.62
    for i, (label, value) in enumerate(zip(labels, values)):
        x = margin_l + (i + 0.19) * plot_w / len(values)
        y = margin_t + plot_h - value / max_val * plot_h
        draw.rectangle((x, y, x + bar_w, margin_t + plot_h), fill=(39, 111, 191))
        draw.text((x + 2, y - 25), f"{value:.2f}", fill=(30, 30, 30), font=small)
        draw.text((x - 2, margin_t + plot_h + 12), label, fill=(55, 55, 55), font=font)
    draw.line((margin_l, margin_t + plot_h, margin_l + plot_w, margin_t + plot_h), fill=(90, 90, 90), width=2)
    img.save(path)


def create_figures(daily_predictions: pd.DataFrame, trend_predictions: pd.DataFrame, trend_metrics: pd.DataFrame) -> dict[str, Path]:
    figs: dict[str, Path] = {}
    sample = daily_predictions.copy()
    if len(sample) > 220:
        sample = sample.iloc[:220].copy()
    labels = pd.to_datetime(sample["date"]).dt.strftime("%Y-%m-%d").tolist()
    figs["daily_time"] = FIG_DIR / "daily_prediction_timeseries.png"
    draw_line_chart(
        figs["daily_time"],
        "Daily PM2.5 Prediction on Test Period",
        labels,
        [
            ("Observed", sample["PM2.5"].tolist(), (24, 44, 78)),
            ("Weather-only", sample["weather_only_pred_PM2.5"].tolist(), (198, 119, 42)),
            ("Best ensemble", sample["ensemble_pred_PM2.5"].tolist(), (39, 111, 191)),
        ],
    )

    figs["daily_scatter"] = FIG_DIR / "best_ensemble_observed_vs_predicted.png"
    draw_scatter_chart(
        figs["daily_scatter"],
        "Best Ensemble: Observed vs Predicted",
        daily_predictions["PM2.5"],
        daily_predictions["ensemble_pred_PM2.5"],
    )

    ridge = trend_metrics[trend_metrics["model"].str.startswith("Trend optimized ensemble")].sort_values("horizon")
    figs["trend_rmse"] = FIG_DIR / "trend_horizon_rmse.png"
    draw_bar_chart(figs["trend_rmse"], "Trend Forecast RMSE by Horizon", [f"H{h}" for h in ridge["horizon"]], ridge["RMSE"].tolist())

    anchors = trend_predictions.groupby("anchor_date")["PM2.5"].mean().sort_values(ascending=False)
    anchor = anchors.index[0]
    trajectory = trend_predictions[trend_predictions["anchor_date"] == anchor].sort_values("horizon")
    labels = pd.to_datetime(trajectory["target_date"]).dt.strftime("%m-%d").tolist()
    figs["trend_example"] = FIG_DIR / "trend_example_trajectory.png"
    draw_line_chart(
        figs["trend_example"],
        f"Seven-Day PM2.5 Trend Forecast from {pd.to_datetime(anchor).strftime('%Y-%m-%d')}",
        labels,
        [
            ("Observed", trajectory["PM2.5"].tolist(), (24, 44, 78)),
            ("Predicted", trajectory["predicted_PM2.5"].tolist(), (39, 111, 191)),
            ("Persistence", trajectory["baseline_PM2.5"].tolist(), (178, 73, 47)),
        ],
        width=1100,
        height=700,
    )
    return figs


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
                    run.font.size = Pt(9.5)
                paragraph.paragraph_format.space_after = Pt(2)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table_from_frame(doc: Document, frame: pd.DataFrame, widths: list[int], max_rows: int | None = None) -> None:
    show = frame if max_rows is None else frame.head(max_rows)
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(show.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(show.columns):
            value = row[col]
            text = format_doc_value(value, str(col))
            cells[idx].text = text
    style_table(table, widths)


def format_doc_value(value, column_name: str) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except TypeError:
        pass
    if column_name in {"样本量", "预测步长"}:
        try:
            return str(int(round(float(value))))
        except (TypeError, ValueError):
            return str(value)
    if column_name == "λ":
        try:
            value_float = float(value)
        except (TypeError, ValueError):
            return str(value)
        if not np.isfinite(value_float):
            return ""
        if abs(value_float - round(value_float)) < 1e-9:
            return str(int(round(value_float)))
        return f"{value_float:.2f}"
    if isinstance(value, (float, np.floating)):
        return f"{float(value):.4f}" if abs(float(value)) < 1000 else f"{float(value):.2f}"
    return str(value)


def apply_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_doc_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("PM2.5 预测模型补充实验报告")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("基于前一日信息与连续时间窗口的日度 PM2.5 预测验证")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")


def make_docx_report(
    daily_metrics: pd.DataFrame,
    class_metrics: pd.DataFrame,
    trend_metrics: pd.DataFrame,
    coeffs: pd.DataFrame,
    stat_tables: dict[str, pd.DataFrame],
    figs: dict[str, Path],
    q33: float,
    q67: float,
) -> None:
    doc = Document()
    apply_doc_styles(doc)
    add_doc_title(doc)

    doc.add_paragraph(
        "本补充实验专门服务于论文中的预测建模部分。与前文解释模型不同，本实验严格按时间顺序切分数据，"
        "只允许模型使用预测日前已经能够获得的信息，避免使用当天气象或当天污染物去预测当天 PM2.5 所造成的信息泄漏。"
    )

    doc.add_heading("1. 预测任务与需求调整", level=1)
    doc.add_paragraph(
        "用户提出的第一个任务是根据前一天气象数据预测当天 PM2.5 情况。严格理解时，模型只能使用昨日温度、湿度、风速、降水、风向等气象变量。"
        "但 PM2.5 具有明显时间持续性，单靠气象变量通常不能充分解释污染物浓度。因此本实验同时报告两个模型："
        "一是严格的前一日气象模型，二是加入昨日 PM2.5、昨日气态污染物和滚动污染特征的增强模型。这样既保留原始问题，又提供更有实际预测价值的方案。"
    )
    doc.add_paragraph(
        "第二个任务是根据一段时间的情况预测 PM2.5 走势。本实验将“一段时间”定义为预测日前连续 7 天窗口，"
        "使用该窗口中的 PM2.5、气象变量和污染物变化，直接预测未来 1 到 7 天的 PM2.5 浓度走势。"
    )

    doc.add_heading("2. 数据与验证设计", level=1)
    doc.add_paragraph(
        f"数据来源为现有 analysis_dataset.csv，时间范围为 2022-05-21 至 2026-05-21。"
        f"模型训练与调参使用 2025-05-22 以前的数据，最终测试使用 2025-05-22 至 2026-05-21 的完整留出期。"
        f"PM2.5 情况分类沿用三分位思想：Low/Medium 分界约为 {q33:.2f}，Medium/High 分界约为 {q67:.2f}。"
    )
    doc.add_paragraph(
        "这种验证方式比随机划分更适合预测问题，因为真实预测时未来数据不可见。若随机打乱样本，模型可能借助相邻日期的相似性获得偏高表现。"
    )

    doc.add_heading("3. 模型方法", level=1)
    doc.add_paragraph(
        "本实验采用多模型比较。第一类是标准化后的岭回归，目标函数为：min Σ(yt - Xtβ)^2 + λΣβj^2，"
        "用于给出可解释的系数结构。第二类是自实现随机森林，用于捕捉滞后变量和滚动变量之间的非线性关系。"
        "第三类是历史相似日模型，即在训练集中寻找气象和污染历史最接近的日期，并用相似日期的 PM2.5 加权平均作为预测。"
        "最后使用验证集选择集成权重，将岭回归、随机森林、相似日模型和持续性基准组合，以尽量降低测试期误差。"
    )

    doc.add_heading("4. 实验链条与不可替代性", level=1)
    chain_rows = pd.DataFrame(
        [
            ["基准层", "训练均值、持续性基准", "平行参照", "给出最低可接受表现，判断新模型是否真的超过“均值”或“今日等于昨日”。"],
            ["气象层", "前日气象模型", "递进起点", "严格回答原始问题，检验昨日温度、湿度、风速、降水、风向本身的预测上限。"],
            ["污染记忆层", "紧凑增强岭回归", "在气象层上递进", "加入昨日 PM2.5、气态污染物和滚动统计量，刻画 PM2.5 的持续性与共同污染。"],
            ["稳健性层", "全滞后岭回归", "与紧凑模型平行", "检查更多 1 至 7 日滞后变量是否带来增益，防止只依赖少数人为挑选特征。"],
            ["非线性层", "滞后随机森林", "与岭回归平行", "检验变量之间是否存在阈值效应和非线性组合，例如高湿低风速下污染累积。"],
            ["相似历史层", "相似日模型", "与参数模型平行", "把历史中相似天气和污染背景的日期作为经验参照，补充回归模型的结构假设不足。"],
            ["综合层", "验证集加权集成", "汇总前述模型", "用验证集自动确定权重，避免主观选择单一模型，并尽量降低留出测试期误差。"],
            ["扩展层", "7 日窗口走势预测", "由单日预测延伸", "回答“未来走势”问题，检验短期窗口能否外推未来 1 至 7 日 PM2.5。"],
        ],
        columns=["层次", "对应实验", "关系", "不可替代性"],
    )
    add_table_from_frame(doc, chain_rows, [1100, 1800, 1300, 3500])
    doc.add_paragraph(
        "上述结构可以理解为：先用基准层确定参照，再用气象层回答原始问题，随后加入污染记忆层提高预测能力，"
        "再通过全滞后、非线性和相似日模型做平行验证，最后用集成模型汇总有效信息；走势预测则是在单日预测基础上的时间外推。"
    )

    doc.add_heading("5. 三类概率统计准确性判定", level=1)
    doc.add_paragraph(
        "正文只保留三类互补的统计判定：第一类用 Bootstrap 置信区间判断误差估计是否稳定；"
        "第二类用配对误差检验判断增强模型是否显著优于基准；第三类用 Conformal 预测区间覆盖率判断预测不确定性是否可控。"
        "其余 MAE、F1、混淆矩阵、残差偏差等结果保留在附录和 results 文件夹中。"
    )

    main_models = [
        "Weather-only previous-day ridge",
        "Compact enhanced ridge",
        "Validation-weighted ensemble",
        "Persistence baseline",
    ]
    ci_show = stat_tables["bootstrap_ci"][stat_tables["bootstrap_ci"]["model"].isin(main_models)].copy()
    ci_show["模型"] = ci_show["model"].map(short_model_name)
    ci_show["RMSE及95%CI"] = ci_show.apply(lambda r: f"{r['RMSE']:.2f} [{r['RMSE_ci_low']:.2f}, {r['RMSE_ci_high']:.2f}]", axis=1)
    ci_show["MAE及95%CI"] = ci_show.apply(lambda r: f"{r['MAE']:.2f} [{r['MAE_ci_low']:.2f}, {r['MAE_ci_high']:.2f}]", axis=1)
    ci_show["R²及95%CI"] = ci_show.apply(lambda r: f"{r['R2']:.3f} [{r['R2_ci_low']:.3f}, {r['R2_ci_high']:.3f}]", axis=1)
    add_table_from_frame(doc, ci_show[["模型", "RMSE及95%CI", "MAE及95%CI", "R²及95%CI"]], [1900, 1900, 1900, 1900])
    doc.add_paragraph("判定一：若 RMSE/MAE 的置信区间整体低于基准模型，说明误差下降不是单个测试样本偶然造成的。")

    paired_show = stat_tables["paired_tests"].head(4).copy()
    paired_show["模型"] = paired_show["model"].map(short_model_name)
    paired_show["参照"] = paired_show["reference"].map(short_model_name)
    paired_show["RMSE降低"] = paired_show["RMSE_reduction_percent"].map(lambda v: f"{v:.2f}%")
    paired_show["p值"] = paired_show["p_value_one_sided"].map(lambda v: f"{v:.4g}")
    add_table_from_frame(doc, paired_show[["模型", "参照", "RMSE降低", "p值", "conclusion"]].rename(columns={"conclusion": "结论"}), [1600, 1600, 1100, 900, 2100])
    doc.add_paragraph("判定二：配对检验逐日比较两个模型的平方误差，能回答“增强模型是否在同一批日期上显著更准”。")

    conformal_show = stat_tables["conformal_intervals"][
        stat_tables["conformal_intervals"]["model"].isin(["Compact enhanced ridge", "Validation-weighted ensemble"])
    ].copy()
    conformal_show["模型"] = conformal_show["model"].map(short_model_name)
    conformal_show["目标覆盖率"] = conformal_show["target_coverage"].map(lambda v: f"{v:.0%}")
    conformal_show["实际覆盖率"] = conformal_show["empirical_coverage"].map(lambda v: f"{v:.2%}")
    conformal_show["区间半宽"] = conformal_show["interval_half_width"]
    conformal_show["覆盖天数"] = conformal_show.apply(lambda r: f"{int(r['covered_days'])}/{int(r['n'])}", axis=1)
    add_table_from_frame(doc, conformal_show[["模型", "目标覆盖率", "实际覆盖率", "区间半宽", "覆盖天数"]], [1900, 1200, 1200, 1200, 1200])
    doc.add_paragraph("判定三：预测区间把点预测扩展为概率判断，若 90% 区间的实际覆盖率接近 90%，说明模型的不确定性估计具有可用性。")

    doc.add_heading("6. 当天 PM2.5 预测展示", level=1)
    ensemble_row = daily_metrics[daily_metrics["model"] == "Validation-weighted ensemble"].iloc[0]
    doc.add_paragraph(
        "验证集选择的最终集成权重为：紧凑增强岭回归 "
        f"{ensemble_row['w_compact']:.2f}，全滞后岭回归 {ensemble_row['w_all_lag']:.2f}，"
        f"滞后随机森林 {ensemble_row['w_forest']:.2f}，相似日模型 {ensemble_row['w_analog']:.2f}，"
        f"持续性基准 {ensemble_row['w_persistence']:.2f}。这说明当前数据中最稳定的主信息来自紧凑增强特征，"
        "随机森林和持续性项提供少量互补修正。"
    )
    doc.add_paragraph(
        "结果显示，严格的前一日气象模型可以提供一定预测信息，但解释力有限；加入昨日污染水平、滚动污染特征和非线性/相似日模型后，"
        "误差整体下降。最终集成模型的权重由验证集自动选择，并与“今日等于昨日”的持续性基准比较，以证明历史气象和污染物信息确实带来了额外预测价值。"
    )
    doc.add_picture(str(figs["daily_time"]), width=Inches(6.3))
    doc.add_paragraph("图1 当天 PM2.5 测试期预测曲线：观测值、严格气象模型和增强模型对比。")
    doc.add_picture(str(figs["daily_scatter"]), width=Inches(5.7))
    doc.add_paragraph("图2 增强模型观测值与预测值散点图。点越接近 1:1 线，说明预测越准确。")

    doc.add_page_break()
    doc.add_heading("7. 未来 1 至 7 日走势预测", level=1)
    doc.add_paragraph(
        "走势模型使用预测日前 7 天窗口作为输入，并分别训练未来第 1 至第 7 天的直接预测模型。"
        "正文用未来各步长的 RMSE 图展示误差随预测步长的变化，完整 MAE、R² 和方向准确率见附录。"
        "一般而言，预测步长越长，误差越高，这是因为未来气象和排放变化的不确定性会不断累积。"
    )
    doc.add_picture(str(figs["trend_rmse"]), width=Inches(6.1))
    doc.add_paragraph("图3 未来 1 至 7 日走势预测 RMSE。")
    doc.add_picture(str(figs["trend_example"]), width=Inches(6.0))
    doc.add_paragraph("图4 测试期某一高污染窗口的 7 日 PM2.5 走势预测示例。")

    doc.add_heading("8. 主要变量与论文解释", level=1)
    top_enhanced = coeffs[coeffs["model"] == "Compact enhanced ridge"].head(10)[
        ["feature", "standardized_coefficient", "abs_standardized_coefficient"]
    ].copy()
    top_enhanced.columns = ["特征", "标准化系数", "绝对值"]
    add_table_from_frame(doc, top_enhanced, [3900, 1700, 1300])
    doc.add_paragraph(
        "紧凑增强模型中的重要特征集中在昨日 PM2.5、近期 PM2.5 均值、NO2、CO、PM10 以及风速、湿度等变量上。"
        "这与前文解释性实验的结论一致：PM2.5 变化同时受到污染持续性、污染物共同变化和气象扩散条件影响。"
    )

    doc.add_heading("9. 后续补充内容", level=1)
    supplement_rows = pd.DataFrame(
        [
            ["未来气象预报", "未来 1 至 7 日温度、湿度、风速、风向、降水预报", "走势预测不能只依赖历史窗口，否则步长越长越缺少未来扩散条件。"],
            ["边界层与稳定度", "边界层高度、气压、逆温或大气稳定度指标", "这些变量直接影响污染物垂直扩散，是高污染峰值误差的重要来源。"],
            ["区域传输", "上风向城市 PM2.5/PM10、风矢量输送指数", "PM2.5 不只来自本地，区域输入会造成单站历史变量无法解释的突增。"],
            ["排放活动", "交通强度、工业排放、节假日或工作日变量", "补充源强变化，使模型能够区分气象导致的累积和排放端增加。"],
            ["概率预测", "预测区间、High 类概率、分类阈值优化", "论文若强调风险预警，除点预测外还应给出高污染概率和不确定性范围。"],
        ],
        columns=["补充方向", "建议变量或公式", "作用"],
    )
    add_table_from_frame(doc, supplement_rows, [1700, 2600, 3100])
    doc.add_paragraph(
        "若写成模型公式，可将当前的点预测扩展为 PM2.5(t+h)=f(历史污染窗口、历史气象窗口、未来气象预报、区域传输、排放活动)+误差项；"
        "其中 h=1,...,7。这样能够把现有统计预测模型与更完整的空气质量预报框架衔接起来。"
    )

    doc.add_heading("10. 可放入论文的结论", level=1)
    doc.add_paragraph(
        "预测实验表明，若只使用前一日气象信息，可以对 PM2.5 情况形成基础判断，但预测能力有限；"
        "加入昨日 PM2.5、气态污染物和 7 日滚动特征后，模型表现更稳定，说明污染持续性和污染物共同变化是 PM2.5 预测中的必要信息。"
        "对于未来走势预测，7 日窗口模型能够给出短期趋势，但随着预测步长增加，误差逐渐扩大。因此，论文中应将该模型定位为短期统计预测模型，"
        "而不是长期空气质量预报系统。若未来希望进一步提高准确率，需要补充未来气象预报、边界层高度、气压、排放清单和区域传输指标。"
    )

    doc.add_page_break()
    doc.add_heading("附录：保留的其他评价结果", level=1)
    doc.add_paragraph(
        "为保持正文集中，以下结果不作为正文的三类核心判定，但全部保留，便于论文答辩或复核时追溯。"
        "对应的完整 CSV 文件同步保存在 results 文件夹中。"
    )
    metric_show = daily_metrics[["model", "n", "RMSE", "MAE", "R2", "Bias", "alpha"]].copy()
    metric_show["model"] = metric_show["model"].map(short_model_name)
    metric_show.columns = ["模型", "样本量", "RMSE", "MAE", "R²", "平均偏差", "λ"]
    add_table_from_frame(doc, metric_show, [2700, 900, 900, 900, 900, 1000, 800])

    cls_show = class_metrics[["model", "accuracy_3class", "weighted_f1_3class", "high_precision", "high_recall", "high_f1"]].copy()
    cls_show["model"] = cls_show["model"].map(short_model_name)
    cls_show.columns = ["模型", "三分类准确率", "加权 F1", "High 精确率", "High 召回率", "High F1"]
    add_table_from_frame(doc, cls_show, [2700, 1200, 1000, 1100, 1100, 1000])

    trend_show = trend_metrics[trend_metrics["model"].str.startswith("Trend direct ridge")][
        ["horizon", "n", "RMSE", "MAE", "R2", "direction_accuracy_vs_anchor", "alpha"]
    ].copy()
    trend_show.columns = ["预测步长", "样本量", "RMSE", "MAE", "R²", "方向准确率", "λ"]
    add_table_from_frame(doc, trend_show, [1000, 900, 900, 900, 900, 1200, 800])

    bias_show = stat_tables["residual_tests"][["model", "bias_mean", "bias_ci_low", "bias_ci_high", "p_value_two_sided", "conclusion"]].copy()
    bias_show["model"] = bias_show["model"].map(short_model_name)
    bias_show["偏差95%CI"] = bias_show.apply(lambda r: f"[{r['bias_ci_low']:.2f}, {r['bias_ci_high']:.2f}]", axis=1)
    bias_show["p值"] = bias_show["p_value_two_sided"].map(lambda v: f"{v:.4g}")
    bias_show = bias_show[["model", "bias_mean", "偏差95%CI", "p值", "conclusion"]]
    bias_show.columns = ["模型", "平均偏差", "偏差95%CI", "p值", "结论"]
    add_table_from_frame(doc, bias_show, [2200, 1200, 1600, 900, 1900])

    acc_ci_show = stat_tables["classification_ci"][["model", "accuracy_3class", "accuracy_ci_low", "accuracy_ci_high", "correct_days", "n"]].copy()
    acc_ci_show["model"] = acc_ci_show["model"].map(short_model_name)
    acc_ci_show["三分类准确率95%CI"] = acc_ci_show.apply(
        lambda r: f"{r['accuracy_3class']:.3f} [{r['accuracy_ci_low']:.3f}, {r['accuracy_ci_high']:.3f}]",
        axis=1,
    )
    acc_ci_show["正确天数"] = acc_ci_show.apply(lambda r: f"{int(r['correct_days'])}/{int(r['n'])}", axis=1)
    acc_ci_show = acc_ci_show[["model", "三分类准确率95%CI", "正确天数"]]
    acc_ci_show.columns = ["模型", "三分类准确率95%CI", "正确天数"]
    add_table_from_frame(doc, acc_ci_show, [2700, 2600, 1300])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("PM2.5 prediction supplement")

    doc.save(DOCX_PATH)


def short_model_name(name: str) -> str:
    mapping = {
        "Weather-only previous-day ridge": "前日气象模型",
        "Compact enhanced ridge": "紧凑增强岭回归",
        "All-lag enhanced ridge": "全滞后岭回归",
        "Enhanced lag random forest": "滞后随机森林",
        "Historical analog KNN": "相似日模型",
        "Validation-weighted ensemble": "验证集加权集成",
        "Persistence baseline": "持续性基准",
        "Train-mean baseline": "训练均值基准",
    }
    return mapping.get(name, name)


def write_readme() -> None:
    readme = ROOT / "README.md"
    readme.write_text(
        "\n".join(
            [
                "# PM2.5 Prediction Model Supplement",
                "",
                "This folder contains reproducible prediction experiments for the PM2.5 paper.",
                "",
                "## Main command",
                "",
                "```bash",
                "'/Users/chenlichong/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3' run_prediction_models.py",
                "```",
                "",
                "## Outputs",
                "",
                "- `data/prediction_feature_dataset.csv`: daily lag feature dataset.",
                "- `data/daily_forecast_test_predictions.csv`: test predictions for same-day PM2.5 from previous-day information.",
                "- `data/trend_anchor_feature_dataset.csv`: 7-day historical-window feature dataset.",
                "- `data/trend_forecast_test_predictions.csv`: 1-7 day trend predictions.",
                "- `results/*.csv`: metrics, validation logs, coefficients, and classification tables.",
                "- `results/stat_*.csv`: bootstrap confidence intervals, paired error tests, residual bias tests, conformal intervals, and Wilson accuracy intervals.",
                "- `figures/*.png`: charts used in the Word report.",
                "- `PM25_prediction_model_report.docx`: paper-ready prediction-model section.",
                "",
                "The split is chronological: training/validation before 2025-05-22, final test from 2025-05-22 to 2026-05-21.",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    ensure_dirs()
    source = read_source()
    daily_data, strict_features, enhanced_features = build_daily_dataset(source)
    daily_predictions, daily_metrics, class_metrics, coeffs, stat_tables, q33, q67 = write_daily_outputs(daily_data, strict_features, enhanced_features)
    trend_predictions, trend_metrics, trend_coeffs = write_trend_outputs(source, q33, q67)
    all_metrics = pd.concat([daily_metrics, trend_metrics], ignore_index=True, sort=False)
    all_metrics.to_csv(RESULTS_DIR / "all_model_metrics.csv", index=False)
    figs = create_figures(daily_predictions, trend_predictions, trend_metrics)
    make_docx_report(daily_metrics, class_metrics, trend_metrics, coeffs, stat_tables, figs, q33, q67)
    write_readme()
    print("Prediction modeling complete.")
    print(f"Daily test metrics: {RESULTS_DIR / 'daily_model_metrics.csv'}")
    print(f"Trend metrics: {RESULTS_DIR / 'trend_horizon_metrics.csv'}")
    print(f"Report: {DOCX_PATH}")


if __name__ == "__main__":
    main()
